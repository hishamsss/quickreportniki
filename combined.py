import streamlit as st
import pandas as pd
import pdfplumber
import re
from docx import Document
from io import BytesIO
from docx.enum.text import WD_COLOR_INDEX

# === Helper Functions ===

def classify(percentile):
    try:
        if isinstance(percentile, str) and ">" in percentile:
            percentile = float(percentile.replace(">", ""))
        elif percentile == "-" or pd.isna(percentile):
            return "-"
        else:
            percentile = float(percentile)
    except:
        return "-"

    if percentile <= 2:
        return "Extremely Low"
    elif 3 <= percentile <= 8:
        return "Borderline"
    elif 9 <= percentile <= 15:
        return "Below Average"
    elif 16 <= percentile <= 24:
        return "Low Average"
    elif 25 <= percentile <= 75:
        return "Average"
    elif 76 <= percentile <= 91:
        return "High Average" 
    elif 92 <= percentile <= 97:
        return "Superior"
    elif percentile >= 98:
        return "Very Superior"
    else:
        return "-"

def format_percentile_with_suffix(percentile):
    try:
        if isinstance(percentile, str) and ">" in percentile:
            percentile = float(percentile.replace(">", ""))
        elif percentile == "-" or pd.isna(percentile):
            return "-"
        else:
            percentile = float(percentile)
    except:
        return "-"

    if percentile.is_integer():
        integer_part = int(percentile)
    else:
        decimal_first_digit = int(str(percentile).split(".")[1][0])
        integer_part = decimal_first_digit

    if 10 <= integer_part % 100 <= 20:
        suffix = 'th'
    else:
        last_digit = integer_part % 10
        if last_digit == 1:
            suffix = 'st'
        elif last_digit == 2:
            suffix = 'nd'
        elif last_digit == 3:
            suffix = 'rd'
        else:
            suffix = 'th'

    if percentile.is_integer():
        return f"{int(percentile)}{suffix}"
    else:
        return f"{percentile}{suffix}"

def replace_placeholders(doc, lookup):
    def copy_run_style(source_run, target_run):
        target_run.font.bold = source_run.font.bold
        target_run.font.italic = source_run.font.italic
        target_run.font.underline = source_run.font.underline
        target_run.font.size = source_run.font.size
        target_run.font.name = source_run.font.name
        if source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb

    def replace_in_runs(runs, lookup):
        pattern = re.compile(r"{{(.*?)}}")
        def _norm_key(k: str) -> str:
            return re.sub(r"\s+", " ", k.strip())
        search_start = 0
        while True:
            full_text = "".join(run.text for run in runs)
            match = pattern.search(full_text, search_start)
            if not match:
                break
            key = _norm_key(match.group(1))
            if key not in lookup:
                search_start = match.end()
                continue
            replacement = lookup[key]
            start, end = match.span()

            # Identify runs affected by the placeholder
            affected_runs = []
            current = 0
            for run in runs:
                run_len = len(run.text)
                run_start = current
                run_end = current + run_len
                if run_end > start and run_start < end:
                    affected_runs.append((run, run_start, run_end))
                if run_end >= end:
                    break
                current = run_end

            if not affected_runs:
                search_start = match.end()
                continue

            start_run, start_run_start, _ = affected_runs[0]
            end_run, end_run_start, _ = affected_runs[-1]
            start_offset = start - start_run_start
            end_offset = end - end_run_start

            if start_run is end_run:
                start_run.text = (
                    start_run.text[:start_offset] + replacement + start_run.text[end_offset:]
                )
            else:
                prefix = start_run.text[:start_offset]
                suffix = end_run.text[end_offset:]
                style_run = affected_runs[1][0] if len(affected_runs) > 1 else start_run
                start_run.text = prefix + replacement
                copy_run_style(style_run, start_run)

                in_between = False
                for run in runs:
                    if in_between:
                        if run is end_run:
                            break
                        run.text = ""
                    if run is start_run:
                        in_between = True

                end_run.text = suffix

            search_start = start + len(replacement)

    for para in doc.paragraphs:
        replace_in_runs(para.runs, lookup)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para.runs, lookup)

def superscript_suffixes(doc):
    pattern = re.compile(r'(\d+(?:\.\d+)?)(st|nd|rd|th)')

    def copy_font_settings(source_run, target_run):
        target_run.font.bold = source_run.font.bold
        target_run.font.italic = source_run.font.italic
        target_run.font.underline = source_run.font.underline
        target_run.font.size = source_run.font.size
        target_run.font.name = source_run.font.name
        if source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb

    def process_runs(paragraph):
        new_runs = []
        for run in paragraph.runs:
            text = run.text
            last_end = 0
            matches = list(pattern.finditer(text))

            if not matches:
                new_runs.append((text, False, run))
            else:
                for match in matches:
                    start, end = match.span()
                    if start > last_end:
                        new_runs.append((text[last_end:start], False, run))
                    new_runs.append((match.group(1), False, run))
                    new_runs.append((match.group(2), True, run))
                    last_end = end
                if last_end < len(text):
                    new_runs.append((text[last_end:], False, run))

        for run in paragraph.runs:
            run.text = ''

        for text, is_super, original_run in new_runs:
            if text == '':
                continue
            new_run = paragraph.add_run(text)
            copy_font_settings(original_run, new_run)
            if is_super:
                new_run.font.superscript = True

    for para in doc.paragraphs:
        process_runs(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_runs(para)

def delete_rows_with_dash(doc):
    for table in doc.tables:
        rows_to_delete = []
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                if cell.text.strip() == "#":
                    rows_to_delete.append(row_idx)
                    break
        for row_idx in sorted(rows_to_delete, reverse=True):
            tbl = table._tbl
            tr = table.rows[row_idx]._tr
            tbl.remove(tr)

def delete_rows_with_unfilled_placeholders(doc):
    pattern = re.compile(r"\{\{.*?\}\}")
    for table in doc.tables:
        rows_to_delete = []
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                if pattern.search(cell.text):
                    rows_to_delete.append(row_idx)
                    break
        for row_idx in sorted(rows_to_delete, reverse=True):
            tbl = table._tbl
            tr = table.rows[row_idx]._tr
            tbl.remove(tr)

def highlight_unfilled_placeholders(doc):
    placeholder_pattern = re.compile(r"\{\{.*?\}\}")
    missing_symbol_pattern = re.compile(r"#")

    def highlight_placeholder_in_runs(runs):
        combined_text = ''
        run_indices = []

        for idx, run in enumerate(runs):
            combined_text += run.text
            run_indices.append(idx)

        matches = list(placeholder_pattern.finditer(combined_text)) + list(missing_symbol_pattern.finditer(combined_text))
        if not matches:
            return

        current_pos = 0
        for idx in run_indices:
            run = runs[idx]
            text_len = len(run.text)
            run_end_pos = current_pos + text_len

            for match in matches:
                match_start, match_end = match.span()
                if match_start < run_end_pos and match_end > current_pos:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

            current_pos = run_end_pos

    for para in doc.paragraphs:
        highlight_placeholder_in_runs(para.runs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    highlight_placeholder_in_runs(para.runs)


def format_cefi_scale_list(df, classification, scale_name_map=None):
    """
    df: CEFI parent DataFrame with columns ['Scale', 'Percentile', 'Percentile*', 'Classification', 'SW']
    classification: e.g., 'Average', 'Superior', etc.
    """
    subset = df[df["Classification"] == classification].copy()
    if subset.empty:
        return ""

    scale_name_map = scale_name_map or {}
    items = []

    for _, row in subset.iterrows():
        scale = row["Scale"]
        display_name = scale_name_map.get(scale, scale)
        items.append(f"{display_name} ({row['Percentile*']} percentile)")

    if len(items) == 1:
        return items[0]
    elif len(items) == 2:
        return " and ".join(items)
    else:
        return "; ".join(items[:-1]) + f"; and {items[-1]}"

def build_cefi_parent_narrative(child_name: str, rater_relation: str = "mother") -> str:
    """
    Build a CEFI narrative from the parent form only, respecting all
    classification levels produced by classify():
    Extremely Low, Borderline, Below Average, Low Average,
    Average, High Average, Superior, Very Superior.
    """
    cefi_df = st.session_state.get("cefi_df")
    if cefi_df is None or cefi_df.empty:
        return ""

    cefi_df = cefi_df.copy()

    # Map internal scale names -> nicer narrative labels (adjust as needed)
    scale_name_map = {
        "Total": "Full scale",
        "Attention": "Attention",
        "Flexibility": "Flexibility",
        "Planning": "Planning",
        "Working Memory": "Working Memory",
        "Emotion Regulation": "Emotion Regulation",
        "Inhibitory Control": "Inhibitory Control",
        "Initiation": "Initiation",
        "Organization": "Organization",
        "Self-Monitoring": "Self-Monitoring",
        # add/edit to match your actual CEFI output
    }

    # ---- Response style from Total row (SW) ----
    response_style = None
    total_rows = cefi_df[cefi_df["Scale"].str.strip().str.lower() == "total"]
    if not total_rows.empty:
        response_style = (total_rows["SW"].iloc[0] or "").strip()

    # First sentence (Response Style)
    if response_style and response_style.lower().startswith("consistent"):
        first_sentence = (
            f"{child_name}'s {rater_relation} demonstrated a Consistent Response Style "
            f"in her answers without indications of positive or negative bias."
        )
    elif response_style:
        first_sentence = (
            f"{child_name}'s {rater_relation} demonstrated a {response_style} Response Style "
            f"in her answers."
        )
    else:
        first_sentence = (
            f"{child_name}'s {rater_relation} completed the rating form."
        )

    sentences = [first_sentence]

    # ---- Classification order and wording (matches your classify() function) ----
    classification_order = [
        "Extremely Low",
        "Borderline",
        "Below Average",
        "Low Average",
        "Average",
        "High Average",
        "Superior",
        "Very Superior",
    ]

    # You can tweak phrasing per band if you want later
    for cls in classification_order:
        scale_list = format_cefi_scale_list(cefi_df, cls, scale_name_map)
        if not scale_list:
            continue

        # For all bands, same basic template:
        # "She rated Ms. Smith in the <cls> range on the following CEFI scales: ..."
        sentences.append(
            f" She rated {child_name} in the {cls} range on the following CEFI scales: {scale_list}."
        )

    return "".join(sentences)

def _format_caars_scale_list(df: pd.DataFrame, guideline: str) -> str:
    """
    Return something like:
    'Inattention/Executive Dysfunction (T=71) and Hyperactivity (T=74)'
    for a given guideline (e.g., 'Very Elevated', 'Not Elevated').
    """
    if df is None or df.empty:
        return ""

    sub = df[df["Guideline"] == guideline].copy()
    if sub.empty:
        return ""

    items = [f"{row['Scale']} (T={row['T-score']})" for _, row in sub.iterrows()]

    if len(items) == 1:
        return items[0]
    elif len(items) == 2:
        return " and ".join(items)
    else:
        return "; ".join(items[:-1]) + f"; and {items[-1]}"


def build_caars2_narrative(name: str, pronoun_cap: str = "Her") -> str:
    """
    Build a CAARS-2 narrative like:

    Ms. Smith reported Very Elevated scores in Inattention/Executive Dysfunction (T=71) and Hyperactivity (T=74),
    as well as Elevated scores in DSM ADHD Inattentive Symptoms (T=64), DSM ADHD Hyperactive/Impulsive Symptoms (T=68),
    and Total ADHD Symptoms (T=67). Her scores for Impulsivity (T=57), Emotional Dysregulation (T=51),
    and Negative Self-Concept (T=43) were Not Elevated. Her ADHD Index was in the Very High range,
    corresponding to a 98% probability.
    """
    content_df = st.session_state.get("caars_content_df", pd.DataFrame())
    dsm_df = st.session_state.get("caars_dsm_df", pd.DataFrame())
    index_info = st.session_state.get("caars_index", {})

    pronoun_lower = pronoun_cap.lower()

    # --- Very Elevated (Content) ---
    very_elevated_content = _format_caars_scale_list(content_df, "Very Elevated")

    # --- Elevated (DSM) ---
    elevated_dsm = _format_caars_scale_list(dsm_df, "Elevated")

    # --- Not Elevated (Content) ---
    not_elevated_content = _format_caars_scale_list(content_df, "Not Elevated")

    sentences = []

    # Sentence 1: Very Elevated + Elevated
    if very_elevated_content or elevated_dsm:
        s1 = f"{name} reported"

        if very_elevated_content:
            s1 += f" Very Elevated scores in {very_elevated_content}"
        if elevated_dsm:
            if very_elevated_content:
                s1 += f", as well as Elevated scores in {elevated_dsm}"
            else:
                s1 += f" Elevated scores in {elevated_dsm}"

        s1 += "."
        sentences.append(s1)

    # Sentence 2: Not Elevated content scales
    if not_elevated_content:
        s2 = f" {pronoun_cap} scores for {not_elevated_content} were Not Elevated."
        sentences.append(s2)

    # Sentence 3: ADHD Index
    idx_guideline = index_info.get("Guideline")
    idx_prob = index_info.get("Probability")
    if idx_guideline and idx_prob:
        s3 = f" {pronoun_cap} ADHD Index was in the {idx_guideline} range, corresponding to a {idx_prob} probability."
        sentences.append(s3)

    return "".join(sentences).strip()


# === Streamlit App ===

st.title("\U0001F4C4 Report Writer")

tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(["WIAT", "Beery", "CEFI", "CVLT", "CAARS-2", "Finalize"])

with tab1:
    uploaded_doc = st.file_uploader("\U0001F4C4 Upload WIAT-4 Report (.docx)", type="docx", key="wiat_upload")

with tab2:
    st.subheader("✍️ Enter Beery Scores")

    col1, col2 = st.columns(2)
    with col1:
        vmi_raw = st.text_input("VMI Raw Score", key="vmi_raw_input")
    with col2:
        vmi = st.text_input("Visual-Motor Integration (VMI) Percentile", key="vmi_input")

    col1, col2 = st.columns(2)
    with col1:
        vp_raw = st.text_input("VP Raw Score", key="vp_raw_input")
    with col2:
        vp = st.text_input("Visual Perception (VP) Percentile", key="vp_input")
        
    col1, col2 = st.columns(2)
    with col1:
        mc_raw = st.text_input("MC Raw Score", key="mc_raw_input")
    with col2:
        mc = st.text_input("Motor Coordination (MC) Percentile", key="mc_input")

with tab3:
    st.subheader("CEFI")

    uploaded_cefi_parent = st.file_uploader(
        "Upload CEFI Parent Report (.pdf)", type="pdf", key="cefi_parent_upload"
    )
    uploaded_cefi_teacher = st.file_uploader(
        "Upload CEFI Teacher Report (.pdf)", type="pdf", key="cefi_teacher_upload"
    )
    
    def _norm_scale(s: str) -> str:
        s = re.sub(r'[^A-Za-z ]', '', str(s))   # letters + spaces only
        s = re.sub(r'\s+', ' ', s).strip()      # collapse spaces
        return s

    def _norm_cvlt_label(s: str) -> str:
        # Letters, numbers and spaces only; collapse whitespace
        s = re.sub(r'[^A-Za-z0-9 ]+', ' ', str(s))
        s = re.sub(r'\s+', ' ', s).strip()
        return s
            
    cefi_df = pd.DataFrame()
    if uploaded_cefi_parent:
        try:
            with pdfplumber.open(uploaded_cefi_parent) as pdf:
                tables = pdf.pages[2].extract_tables()
            df = pd.concat([pd.DataFrame(tbl) for tbl in tables], ignore_index=True)
            valid_row_drops = [i for i in [0, 1, 3, 4] if 0 <= i < len(df)]
            df = df.drop(df.index[valid_row_drops]).reset_index(drop=True)
            if df.shape[1] > 4:
                df.iat[0, 3] = df.iat[0, 4]
                df.iat[0, 4] = ""
            df.iat[0, 0] = "Total"
            valid_col_drops = [c for c in [1, 2, 4, 5, 6] if c in df.columns]
            df = df.drop(columns=valid_col_drops).reset_index(drop=True)
            cefi_df = df.copy()
            cefi_df.columns = ["Scale", "Percentile", "SW"]
            cefi_df["Scale"] = cefi_df["Scale"].apply(_norm_scale)
            cefi_df["SW"] = cefi_df["SW"].replace({"None": "N/A"}).fillna("N/A")
            cefi_df["Classification"] = cefi_df["Percentile"].apply(classify)
            cefi_df["Percentile*"] = cefi_df["Percentile"].apply(format_percentile_with_suffix)
            st.session_state["cefi_df"] = cefi_df
        except Exception as e:
            st.error(f"Error processing CEFI Parent PDF: {e}")
            st.exception(e)

    cefi_teacher_df = pd.DataFrame()
    if uploaded_cefi_teacher:
        try:
            with pdfplumber.open(uploaded_cefi_teacher) as pdf:
                tables = pdf.pages[2].extract_tables()
            df = pd.concat([pd.DataFrame(tbl) for tbl in tables], ignore_index=True)
            valid_row_drops = [i for i in [0, 1, 3, 4] if 0 <= i < len(df)]
            df = df.drop(df.index[valid_row_drops]).reset_index(drop=True)
            if df.shape[1] > 4:
                df.iat[0, 3] = df.iat[0, 4]
                df.iat[0, 4] = ""
            df.iat[0, 0] = "Total"
            valid_col_drops = [c for c in [1, 2, 4, 5, 6] if c in df.columns]
            df = df.drop(columns=valid_col_drops).reset_index(drop=True)
            cefi_teacher_df = df.copy()
            cefi_teacher_df.columns = ["Scale", "Percentile", "SW"]
            cefi_teacher_df["Scale"] = cefi_teacher_df["Scale"].apply(_norm_scale)
            cefi_teacher_df["SW"] = cefi_teacher_df["SW"].replace({"None": "N/A"}).fillna("N/A")
            cefi_teacher_df["Classification"] = cefi_teacher_df["Percentile"].apply(classify)
            cefi_teacher_df["Percentile*"] = cefi_teacher_df["Percentile"].apply(format_percentile_with_suffix)
            st.session_state["cefi_teacher_df"] = cefi_teacher_df
        except Exception as e:
            st.error(f"Error processing CEFI Teacher PDF: {e}")
            st.exception(e)

with tab4:
    st.subheader("CVLT-3")

    uploaded_cvlt = st.file_uploader(
        "Upload CVLT-3 Report (.pdf)",
        type="pdf",
        key="cvlt_upload"
    )

    cvlt_info = {}
    cvlt_scores = pd.DataFrame()

    if uploaded_cvlt:
        try:
            with pdfplumber.open(uploaded_cvlt) as pdf:
                # Only pages 1, 3, 4, 5, 6  -> zero-based indices 0,2,3,4,5
                page_indices = [1, 3, 4, 5, 6]
                texts = []

                for idx in page_indices:
                    if idx < len(pdf.pages):
                        page_text = pdf.pages[idx].extract_text() or ""
                        texts.append(page_text)

                        # --- Page 1: demographic/info fields ---
                        if idx == 0:
                            flat = " ".join(page_text.splitlines())

                            m = re.search(r"ID:\s*([^\s]+)", flat)
                            if m:
                                cvlt_info["ID"] = m.group(1)

                            m = re.search(r"Name:\s*([^G]+)Gender:", flat)
                            if m:
                                cvlt_info["Name"] = m.group(1).strip()

                            m = re.search(r"Test Date:\s*([0-9/]+)", flat)
                            if m:
                                cvlt_info["Test Date"] = m.group(1)

                            m = re.search(r"Examiner Name:\s*([^B]+)Birth Date:", flat)
                            if m:
                                cvlt_info["Examiner Name"] = m.group(1).strip()

                            m = re.search(r"Gender:\s*([A-Za-z]+)", flat)
                            if m:
                                cvlt_info["Gender"] = m.group(1)

                            m = re.search(r"Birth Date:\s*([0-9/]+)", flat)
                            if m:
                                cvlt_info["Birth Date"] = m.group(1)

                            m = re.search(r"Age at Testing:\s*([0-9 ]+years [0-9 ]+months)", flat)
                            if m:
                                cvlt_info["Age at Testing"] = m.group(1)

                full_text = "\n".join(texts)

            # --- Generic score parser for pages 3–6 ---
            rows = []
            for line in full_text.splitlines():
                line = line.strip()
                if not line:
                    continue

                # Lines with: label  number  number  number
                m = re.match(
                    r"^([A-Za-z0-9–%/'(),\. ]+?)\s+(-?\d+\.?\d*)\s+(-?\d+\.?\d*)\s+(-?\d+\.?\d*)$",
                    line
                )
                if m:
                    label, c2, c3, c4 = m.groups()
                    rows.append({
                        "Label": label.strip(),
                        "Col2": c2,
                        "Col3": c3,
                        "Col4": c4
                    })

            if rows:
                cvlt_scores = pd.DataFrame(rows)
                st.dataframe(cvlt_scores, use_container_width=True)

            st.session_state["cvlt_info"] = cvlt_info
            st.session_state["cvlt_scores"] = cvlt_scores

        except Exception as e:
            st.error(f"Error processing CVLT PDF: {e}")
            st.exception(e)

with tab5:  # your CAARS tab
    st.subheader("CAARS-2 Self-Report")

    uploaded_caars = st.file_uploader(
        "Upload CAARS-2 Self-Report (.pdf)",
        type="pdf",
        key="caars_upload",
    )

    if uploaded_caars:
        try:
            with pdfplumber.open(uploaded_caars) as pdf:
                # Page 4 in the report (0-based index 3)
                tables = pdf.pages[3].extract_tables()

            # Helper to find columns by fuzzy name
            import re

            def _find_cols(df, want_t=True, want_guideline=True):
                cols = list(df.columns)
                scale_col = cols[0]  # first column is always the scale name

                t_col = None
                guid_col = None

                for c in cols:
                    name = str(c)
                    if want_t and t_col is None and re.search(r"t.?score", name, re.I):
                        t_col = c
                    if want_guideline and guid_col is None and "guideline" in name.lower():
                        guid_col = c

                return scale_col, t_col, guid_col

            # --- Content Scales table ---
            content_raw = pd.DataFrame(tables[0])
            content_raw.columns = content_raw.iloc[0]
            content_raw = content_raw.drop(index=0).reset_index(drop=True)

            scale_col, t_col, guid_col = _find_cols(content_raw)

            content_df = content_raw[[scale_col, t_col, guid_col]].copy()
            content_df.columns = ["Scale", "T-score", "Guideline"]

            # --- DSM Scales table ---
            dsm_raw = pd.DataFrame(tables[1])
            dsm_raw.columns = dsm_raw.iloc[0]
            dsm_raw = dsm_raw.drop(index=0).reset_index(drop=True)

            scale_col, t_col, guid_col = _find_cols(dsm_raw)

            dsm_df = dsm_raw[[scale_col, t_col, guid_col]].copy()
            dsm_df.columns = ["Scale", "T-score", "Guideline"]

            # --- ADHD Index table ---
            index_raw = pd.DataFrame(tables[2])
            index_raw.columns = index_raw.iloc[0]
            index_raw = index_raw.drop(index=0).reset_index(drop=True)

            # find "Probability" and "Guideline" columns flexibly
            prob_col = None
            idx_guid_col = None
            for c in index_raw.columns:
                name = str(c).lower()
                if prob_col is None and "probab" in name:
                    prob_col = c
                if idx_guid_col is None and "guideline" in name:
                    idx_guid_col = c

            adhd_index_prob = str(index_raw.loc[0, prob_col]).strip() if prob_col else ""
            adhd_index_guideline = str(index_raw.loc[0, idx_guid_col]).strip() if idx_guid_col else ""

            # Store for later
            st.session_state["caars_content_df"] = content_df
            st.session_state["caars_dsm_df"] = dsm_df
            st.session_state["caars_index"] = {
                "Probability": adhd_index_prob,
                "Guideline": adhd_index_guideline,
            }

            # DEBEUG REMOVE
            st.write("Content DF preview:")
            st.dataframe(content_df.head())
            
            st.write("DSM DF preview:")
            st.dataframe(dsm_df.head())
            
            st.write("Index info:")
            st.json(st.session_state["caars_index"])

        except Exception as e:
            st.error(f"Error processing CAARS-2 PDF: {e}")
            st.exception(e)

with tab6:
    st.subheader("Report Settings")

    # 1) Always-visible fields:
    report_name_input = st.text_input(
        "Report file name (without .docx)",
        value="combined_report",
        key="report_name_input"
    )  
    gender_selection = st.radio(
        "Select WIAT Report Gender Template:",
        ("Male", "Female"),
        key="gender"
    )

    # 2) If files aren’t uploaded yet, prompt the user:
    if not uploaded_doc:
        st.info("Please upload both your WIAT report in the WIAT tab.")
    else:
        # 3) Once both are present, show the generate button
        if st.button("Generate Combined Report"):
            # … your existing document-generation logic here …
            input_doc    = Document(uploaded_doc)
            template_path = (
                "n_male_template.docx"
                if gender_selection == "Male"
                else "n_female_template.docx"
            )
            template_doc = Document(template_path)
        
            # 2) (Now your existing AE‐table loops, placeholder‐replacing, superscripting, etc.)
        
            # 3) Save into bytes
            output = BytesIO()
            template_doc.save(output)
            st.session_state["generated_report"] = output.getvalue()            
            st.success("✅ Combined document generated successfully!")

            # === Process WIAT Tables ===
            ae_combined = pd.DataFrame()

            for i, table in enumerate(input_doc.tables):
                data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                df = pd.DataFrame(data)
                if df.shape[0] > 1:
                    df.columns = df.iloc[0]
                    df = df.drop(index=0).reset_index(drop=True)
                if df.shape[1] >= 5:
                    ae_df = df.iloc[:, [0, 4]].copy()
                    ae_df.columns = ['Name', 'Percentile']
                    ae_df['Name'] = ae_df['Name'].str.replace(r'[^A-Za-z\s]', '', regex=True).str.strip()
                    ae_combined = pd.concat([ae_combined, ae_df], ignore_index=True)
                    
            if not ae_combined.empty:
                ae_combined.drop_duplicates(subset='Name', inplace=True)
                ae_combined["Classification"] = ae_combined["Percentile"].apply(classify)
                ae_combined["Percentile*"] = ae_combined["Percentile"].apply(format_percentile_with_suffix)
                ae_combined = ae_combined.replace("-", "#")

            lookup = {}
            for _, row in ae_combined.iterrows():
                name = row['Name'].strip()
                lookup[f"{name} Classification"] = row['Classification']
                lookup[f"{name} Percentile"] = str(row['Percentile']).strip()
                lookup[f"{name} Percentile*"] = str(row['Percentile*']).strip()

            # === Beery
            if vmi:
                lookup["VMI Percentile"] = vmi
                lookup["VMI Percentile*"] = format_percentile_with_suffix(vmi)
                lookup["VMI Classification"] = classify(vmi)
            if vmi_raw:
                lookup["VMI Raw Score"] = vmi_raw
            if vp:
                lookup["VP Percentile"] = vp
                lookup["VP Percentile*"] = format_percentile_with_suffix(vp)
                lookup["VP Classification"] = classify(vp)
            if vp_raw:
                lookup["VP Raw Score"] = vp_raw
            if mc:
                lookup["MC Percentile"] = mc
                lookup["MC Percentile*"] = format_percentile_with_suffix(mc)
                lookup["MC Classification"] = classify(mc)
            if mc_raw:
                lookup["MC Raw Score"] = mc_raw

            # === CEFI Parent
            if not cefi_df.empty:
                for _, row in cefi_df.iterrows():
                    scale = row['Scale'].strip()
                    lookup[f"CEFI {scale} Classification"] = row['Classification']
                    lookup[f"CEFI {scale} Percentile"]      = str(row['Percentile']).strip()
                    lookup[f"CEFI {scale} Percentile*"]     = str(row['Percentile*']).strip()
                    lookup[f"CEFI {scale} SW"]              = str(row['SW']).strip()
                    
            # === CEFI Teacher
            if not cefi_teacher_df.empty:
                for _, row in cefi_teacher_df.iterrows():
                    scale = row['Scale'].strip()
                    lookup[f"CEFI Teacher {scale} Classification"] = row['Classification']
                    lookup[f"CEFI Teacher {scale} Percentile"]      = str(row['Percentile']).strip()
                    lookup[f"CEFI Teacher {scale} Percentile*"]     = str(row['Percentile*']).strip()
                    lookup[f"CEFI Teacher {scale} SW"]              = str(row['SW']).strip()
            
            def _prefill_missing_cefi_channel(lookup: dict, channel: str, scales: set):
                prefix = "CEFI" if channel == "Parent" else "CEFI Teacher"
                for sc in sorted(scales):
                    lookup.setdefault(f"{prefix} {sc} Percentile", "N/A")
                    lookup.setdefault(f"{prefix} {sc} Percentile*", "N/A")
                    lookup.setdefault(f"{prefix} {sc} Classification", "N/A")
                    lookup.setdefault(f"{prefix} {sc} SW", "N/A")

            all_cefi_scales = set()
            if not cefi_df.empty:
                all_cefi_scales |= set(cefi_df["Scale"])
            if not cefi_teacher_df.empty:
                all_cefi_scales |= set(cefi_teacher_df["Scale"])

            if (not cefi_df.empty) and cefi_teacher_df.empty:
                _prefill_missing_cefi_channel(lookup, "Teacher", all_cefi_scales)
            if cefi_df.empty and (not cefi_teacher_df.empty):
                _prefill_missing_cefi_channel(lookup, "Parent", all_cefi_scales)

            if not cefi_df.empty and not cefi_teacher_df.empty:
                lookup["CEFI Heading"] = "The percentiles for the parent and teacher rating scales are presented in the table that follows for comparison."
            elif not cefi_df.empty:
                lookup["CEFI Heading"] = "The percentiles for the parent rating scales are presented in the table that follows."
            elif not cefi_teacher_df.empty:
                lookup["CEFI Heading"] = "The percentiles for the teacher rating scales are presented in the table that follows."

            # === CEFI Parent narrative ===
            if st.session_state.get("cefi_df") is not None and not st.session_state["cefi_df"].empty:
                parent_narrative = build_cefi_parent_narrative(
                    child_name="Ms. Smith",
                    rater_relation="mother"
                )
                lookup["CEFI Parent Narrative"] = parent_narrative


            # === CVLT (from CVLT tab) ===
            cvlt_info = st.session_state.get("cvlt_info", {})
            cvlt_scores = st.session_state.get("cvlt_scores", pd.DataFrame())

            # Demographic / header info from page 1
            for key, value in cvlt_info.items():
                # e.g., "CVLT Name", "CVLT Test Date"
                lookup[f"CVLT {key}"] = str(value).strip()

            # Score tables from pages 3–6
            if isinstance(cvlt_scores, pd.DataFrame) and not cvlt_scores.empty:
                for _, row in cvlt_scores.iterrows():
                    label = row["Label"]
                    norm_label = _norm_cvlt_label(label)

                    # Col2 / Col3 / Col4 are the three numeric columns on that line
                    # (e.g., Sum of scaled scores / Index score / Percentile rank
                    #  in the Standard Score Summary section).
                    lookup[f"CVLT {norm_label} Col2"] = str(row["Col2"]).strip()
                    lookup[f"CVLT {norm_label} Col3"] = str(row["Col3"]).strip()
                    lookup[f"CVLT {norm_label} Percentile"] = str(row["Col4"]).strip()

                    # Optional: if Col4 is a percentile, you can also classify it:
                    try:
                        lookup[f"CVLT {norm_label} Classification"] = classify(row["Col4"])
                        lookup[f"CVLT {norm_label} Percentile*"] = format_percentile_with_suffix(row["Col4"])
                    except Exception:
                        pass

            # === CAARS-2 Narrative ===
            if (
                st.session_state.get("caars_content_df") is not None
                and not st.session_state["caars_content_df"].empty
            ):
                caars_narr = build_caars2_narrative(name="Ms. Smith", pronoun_cap="Her")
                lookup["CAARS2 Narrative"] = caars_narr


            # === Fill and output unified report
            lookup = {re.sub(r"\s+", " ", k.strip()): v for k, v in lookup.items()}
            replace_placeholders(template_doc, lookup)
            superscript_suffixes(template_doc)
            delete_rows_with_dash(template_doc)
            delete_rows_with_unfilled_placeholders(template_doc)
            highlight_unfilled_placeholders(template_doc)

            output = BytesIO()
            template_doc.save(output)
            st.session_state["generated_report"] = output.getvalue()
            

        if st.session_state.get("generated_report"):
            output_data = BytesIO(st.session_state["generated_report"])
            final_name = report_name_input.strip() or "combined_report"
            if not final_name.lower().endswith(".docx"):
                final_name += ".docx"

            st.download_button(
                label="📥 Download Combined Report",
                data=output_data,
                file_name=final_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
